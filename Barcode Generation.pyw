from PyQt5 import QtWidgets, uic
from datetime import *
import sys
import os
from gen_classes import HospitexDB
from configparser import *
import docx


class Ui(QtWidgets.QMainWindow):
    def __init__(self):
        super(Ui, self).__init__()
        uic.loadUi('bars.ui', self)
        self.conf = ConfigParser()
        self.conf.read("Settings.ini", encoding="utf-8")
        self.params = self.conf['Parameters']
        self.device = self.params['device']
        self.hosp = self.params['hospital']
        self.dev_uid = self.params['device_uid']
        self.dev_sn = self.params['device_sn']
        self.is_from_invoice = self.params['task_from_invoice']

        self.dev_type_box.addItems(['Tecom', 'Urit', 'Bioelab'])
        self.dev_type_box.setCurrentText(self.device)
        self.invoices = last_invoices()

        for invoice in reversed(self.invoices):
            self.invoice_box.addItem(*invoice)
        self.invoice_box.setCurrentText("Выберите счет")

        self.doc = docx.Document('Коды приборов.docx')
        self.dev_tables = {'Tecom': 0, 'Urit': 3, 'Bioelab': 1}
        self.get_hosps()

        if self.is_from_invoice == 'No':
            self.manual_rb.setChecked(True)
        elif self.is_from_invoice == 'Yes':
            self.invoice_rb.setChecked(True)
            self.invoice_box.setEnabled(True)

        self.invoice_rb.toggled.connect(self.by_invoice)
        self.manual_rb.toggled.connect(self.from_file)
        self.invoice_box.activated.connect(self.add_btn_enable)
        self.hosp_box.activated.connect(self.ok_btn_enable)
        self.add_btn.clicked.connect(self.add_invoice)
        self.dev_type_box.activated.connect(self.get_hosps)
        self.ok_btn.clicked.connect(self.start)
        self.show()

    def by_invoice(self):
        self.invoice_box.setEnabled(True)

    def from_file(self):
        self.invoice_box.setEnabled(False)
        self.ok_btn.setEnabled(False)
        self.add_btn.setEnabled(False)
        self.invoice_list.clear()
        self.invoice_list.setEnabled(False)

    def add_btn_enable(self):
        self.add_btn.setEnabled(True)

    def ok_btn_enable(self):
        self.ok_btn.setEnabled(True)

    def add_invoice(self):
        self.invoice_list.setEnabled(True)
        self.invoice_list.addItem(self.invoice_box.currentText())

    def start(self):
        if self.invoice_rb.isChecked():
            self.is_from_invoice = 'Yes'
        elif self.manual_rb.isChecked():
            self.is_from_invoice = 'No'
        print(self.is_from_invoice)
        self.params.set('task_from_invoice', self.is_from_invoice)
        print(self.is_from_invoice)
        #
        # self.params.set('invoice_n', f'{self.invoice_list.items().join(",")}')
        # self.params.set('device', self.dev_type_box.currentText())
        #
        # device = self.dev_type_box.currentText()
        # table = self.doc.tables[self.dev_tables[device]]
        # sn = table.cell(self.hosp_box.currentIndex(), 2)
        # uid = table.cell(self.hosp_box.currentIndex(), 3)
        # self.params.set('device_uid', uid)
        # self.params.set('device_sn'. sn)
        # self.params.set('hospital', self.hosp_box.currentText())

        with open("Settings.ini", "w", encoding="utf-8") as config_file:
            self.conf.write(config_file)
        # self.close()

        # os.system(f'python {device.lower()}.py')

    def get_hosps(self):
        device = self.dev_type_box.currentText()
        self.hosp_box.clear()
        print(device)
        table = self.doc.tables[self.dev_tables[device]]
        result = []
        for r in range(1, len(table.rows)):
            result.append(table.cell(r, 0).text)

        self.hosp_box.addItems(result)
        self.hosp_box.setCurrentText("Выберите прибор")


def last_invoices():
    trade_db = HospitexDB("trade")
    now_date = datetime.now()
    days_ago = now_date - timedelta(days=14)

    results = trade_db.db_request("select NOM_SH from EXCEL_SHET "
                                  "where DATA between {d'%(da)s'} and {d'%(nwd)s'}"
                                  % {'da': days_ago.date(),
                                     'nwd': now_date.date()}
                                  )
    return results


app = QtWidgets.QApplication(sys.argv)
window = Ui()
app.exec_()
